from io import BytesIO
from docx import Document
from docx.shared import Pt

def generate_docx_from_template(template_text: str, data: dict) -> BytesIO:
    """
    Создаёт документ Word из текстового шаблона и словаря с данными.
    Заменяет все вхождения {key} на соответствующие значения.
    Возвращает BytesIO с готовым файлом.
    """
    # Подстановка данных
    for key, value in data.items():
        template_text = template_text.replace("{" + key + "}", value)

    # Создаём новый документ
    doc = Document()

    # Разбиваем текст на строки и добавляем их как параграфы
    for line in template_text.split('\n'):
        if line.strip() == '':
            doc.add_paragraph()  # пустая строка
        else:
            p = doc.add_paragraph(line.strip())
            # Базовая настройка шрифта (можно убрать или настроить)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

    # Сохраняем в поток
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream